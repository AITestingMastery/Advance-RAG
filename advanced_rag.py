from __future__ import annotations
import csv, io, json, os, re, sqlite3, uuid
from pathlib import Path
from typing import Any
import chromadb
import numpy as np
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document as DocxDocument
from rank_bm25 import BM25Okapi

load_dotenv()
CHAT_MODEL = os.getenv("RAG_CHAT_MODEL", "gpt-4o-mini")
EMBED_MODEL = os.getenv("RAG_EMBED_MODEL", "text-embedding-3-small")
TOP_K, CHUNK_SIZE, CHUNK_OVERLAP, RRF_K = 5, 700, 100, 60
# Documents larger than this many characters are split into line-aligned
# batches before structured-record extraction, so one LLM call never has to
# swallow an entire book and so we never cut a record row in half.
MAX_STRUCT_CHARS = 40_000
BASE_DIR = Path(__file__).resolve().parent
STORAGE_DIR = BASE_DIR / "storage"; STORAGE_DIR.mkdir(exist_ok=True)
SQLITE_PATH = STORAGE_DIR / "rag.db"
CHROMA_PATH = STORAGE_DIR / "chroma"
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".csv", ".txt", ".md", ".markdown")
_client: OpenAI | None = None
_chroma = chromadb.PersistentClient(path=str(CHROMA_PATH))
_collection = _chroma.get_or_create_collection(name="rag_chunks", metadata={"hnsw:space": "cosine"})


def get_client():
    global _client
    if _client is None:
        key = os.getenv("OPENAI_API_KEY", "")
        if not key.startswith("sk-"): raise RuntimeError("OPENAI_API_KEY is missing.")
        _client = OpenAI(api_key=key)
    return _client


def db():
    conn = sqlite3.connect(SQLITE_PATH); conn.row_factory = sqlite3.Row; return conn


def init_db():
    with db() as conn:
        conn.execute("""CREATE TABLE IF NOT EXISTS documents (
            document_id TEXT PRIMARY KEY,
            filename TEXT NOT NULL,
            file_type TEXT NOT NULL,
            units INTEGER NOT NULL,
            chunks INTEGER NOT NULL,
            records INTEGER NOT NULL
        )""")
        conn.execute("CREATE TABLE IF NOT EXISTS records (id INTEGER PRIMARY KEY AUTOINCREMENT, document_id TEXT NOT NULL, record_json TEXT NOT NULL)")
        conn.commit()
init_db()


# --------------------------------------------------------------------------
# File loaders: one function per format, all normalized to
# (text_for_search, unit_count, structured_records_or_None).
# structured_records is None when the format has no inherent tabular
# structure (PDF / TXT / MD / DOCX-without-tables) — in that case the caller
# falls back to LLM-based extraction on the text. CSV and DOCX tables are
# already structured, so their records are parsed programmatically, which is
# faster, free, and never hallucinates a value.
# --------------------------------------------------------------------------

def read_text_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def extract_pdf(file_bytes: bytes):
    reader = PdfReader(io.BytesIO(file_bytes)); parts = []
    for page_no, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip(): parts.append(f"--- Page {page_no} ---\n{text}")
    text = "\n".join(parts)
    if not text.strip():
        raise ValueError("No readable text found in PDF (it may be a scanned/image-only PDF, which needs OCR — not supported yet).")
    return text, len(reader.pages)


def extract_docx(file_bytes: bytes):
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text_parts = list(paragraphs)
    table_records: list[dict] = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        headers = rows[0]
        for row in rows[1:]:
            record = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
            table_records.append(record)
            text_parts.append(" | ".join(f"{headers[i]}: {row[i]}" for i in range(min(len(headers), len(row)))))
    text = "\n".join(text_parts)
    if not text.strip():
        raise ValueError("No readable text found in DOCX.")
    units = len(paragraphs) + sum(len(t.rows) for t in doc.tables)
    return text, units, (table_records or None)


def extract_csv(file_bytes: bytes):
    raw = read_text_bytes(file_bytes)
    try:
        dialect = csv.Sniffer().sniff(raw[:4096], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    reader = csv.DictReader(io.StringIO(raw), dialect=dialect)
    records, text_parts = [], []
    for row in reader:
        clean = {(k or "").strip(): (v.strip() if isinstance(v, str) else v) for k, v in row.items() if k}
        if not clean:
            continue
        records.append(clean)
        text_parts.append(" | ".join(f"{k}: {v}" for k, v in clean.items()))
    if not records:
        raise ValueError("No rows found in CSV (check it has a header row and at least one data row).")
    return "\n".join(text_parts), records


def extract_plain_text(file_bytes: bytes):
    text = read_text_bytes(file_bytes)
    if not text.strip():
        raise ValueError("File is empty.")
    return text, text.count("\n") + 1


def extract_content(file_bytes: bytes, filename: str):
    """Returns (text, unit_count, structured_records_or_None)."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text, units = extract_pdf(file_bytes)
        return text, units, None
    if ext == ".docx":
        text, units, table_records = extract_docx(file_bytes)
        return text, units, table_records
    if ext == ".csv":
        text, records = extract_csv(file_bytes)
        return text, len(records), records
    if ext in (".txt", ".md", ".markdown"):
        text, units = extract_plain_text(file_bytes)
        return text, units, None
    raise ValueError(f"Unsupported file type '{ext or filename}'. Supported types: {', '.join(SUPPORTED_EXTENSIONS)}")


def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    words = text.split(); chunks = []; current = []; length = 0
    for word in words:
        current.append(word); length += len(word) + 1
        if length >= size:
            chunks.append(" ".join(current))
            back = []; back_len = 0
            for w in reversed(current):
                if back_len >= overlap: break
                back.insert(0, w); back_len += len(w) + 1
            current = back; length = back_len
    if current: chunks.append(" ".join(current))
    return chunks


def tokenize(text): return re.findall(r"[a-z0-9]+", text.lower())


def embed(texts):
    response = get_client().embeddings.create(model=EMBED_MODEL, input=texts)
    return np.array([x.embedding for x in response.data], dtype=np.float32)


def _batch_by_lines(text: str, max_chars: int):
    """Split text into chunks no larger than max_chars, only ever breaking on
    line boundaries so a table row is never split across two LLM calls."""
    if len(text) <= max_chars:
        return [text]
    batches, current, current_len = [], [], 0
    for line in text.split("\n"):
        line_len = len(line) + 1
        if current_len + line_len > max_chars and current:
            batches.append("\n".join(current)); current, current_len = [], 0
        current.append(line); current_len += line_len
    if current: batches.append("\n".join(current))
    return batches


def _extract_records_batch(text: str):
    prompt = f"""Extract repeated structured records from this document.
Return ONLY JSON: {{"is_structured":true,"record_type":"string","records":[]}}
Preserve values exactly. Do not invent. Normalize field names to snake_case.
Include every reliably identifiable record. If not structured, records=[].
DOCUMENT:
{text}"""
    response = get_client().chat.completions.create(model=CHAT_MODEL, temperature=0, response_format={"type": "json_object"}, messages=[
        {"role": "system", "content": "Return valid JSON only."}, {"role": "user", "content": prompt}])
    try:
        data = json.loads(response.choices[0].message.content or "{}")
    except json.JSONDecodeError:
        return []
    records = data.get("records", [])
    return records if isinstance(records, list) else []


def extract_records(text):
    """LLM-based structured extraction, batched over line-aligned windows for
    large documents so we don't exceed context limits or blow up cost."""
    if not text.strip():
        return []
    all_records = []
    for batch in _batch_by_lines(text, MAX_STRUCT_CHARS):
        try:
            all_records.extend(_extract_records_batch(batch))
        except Exception:
            continue
    return all_records


def normalize_record(record):
    out = {}
    for key, value in record.items():
        key = re.sub(r"[^a-z0-9_]+", "_", str(key).strip().lower()).strip("_")
        out[key] = value.strip() if isinstance(value, str) else value
    return out


def process_document(file_bytes: bytes, filename: str):
    ext = Path(filename).suffix.lower().lstrip(".")
    text, units, programmatic_records = extract_content(file_bytes, filename)
    chunks = chunk_text(text)
    if programmatic_records is not None:
        records = [normalize_record(r) for r in programmatic_records]
    else:
        records = [normalize_record(x) for x in extract_records(text)]
    document_id = uuid.uuid4().hex[:12]
    with db() as conn:
        conn.execute("INSERT INTO documents VALUES (?,?,?,?,?,?)", (document_id, filename, ext, units, len(chunks), len(records)))
        for record in records:
            conn.execute("INSERT INTO records(document_id,record_json) VALUES (?,?)", (document_id, json.dumps(record, ensure_ascii=False)))
        conn.commit()
    embeddings = embed(chunks) if chunks else np.empty((0, 0), dtype=np.float32)
    if chunks:
        _collection.add(
            ids=[f"{document_id}_{i}" for i in range(len(chunks))],
            documents=chunks,
            metadatas=[{"document_id": document_id, "filename": filename, "chunk_index": i} for i in range(len(chunks))],
            embeddings=[x.tolist() for x in embeddings])
    return {"document_id": document_id, "filename": filename, "file_type": ext, "units": units, "chunks": len(chunks), "records": len(records)}


def list_documents():
    with db() as conn: rows = conn.execute("SELECT * FROM documents ORDER BY rowid DESC").fetchall()
    return [dict(x) for x in rows]


def delete_document(document_id):
    ids = _collection.get(where={"document_id": document_id}, include=[])["ids"]
    if ids: _collection.delete(ids=ids)
    with db() as conn:
        conn.execute("DELETE FROM records WHERE document_id=?", (document_id,))
        conn.execute("DELETE FROM documents WHERE document_id=?", (document_id,))
        conn.commit()


def delete_all_documents():
    for doc in list_documents():
        delete_document(doc["document_id"])


def get_records(document_id):
    with db() as conn: rows = conn.execute("SELECT record_json FROM records WHERE document_id=?", (document_id,)).fetchall()
    return [json.loads(x["record_json"]) for x in rows]


VALID_STRATEGIES = ("structured_filter", "aggregation", "exact_lookup", "hybrid_retrieval")

# Structured Outputs schema for the router. This is the real fix for a
# failure we observed in the wild: with the older loose {"type":"json_object"}
# mode, the model sometimes echoed the prompt's own placeholder text back as
# the value — e.g. returning "strategy": "structured_filter | aggregation"
# verbatim, which matches none of the strategy branches in query_document()
# and silently falls through to hybrid_retrieval, discarding two perfectly
# good filters that were parsed correctly in the same response. A strict
# JSON Schema with an `enum` makes that response impossible to produce: the
# model is constrained at generation time to one of the four literal values.
_CLASSIFY_SCHEMA = {
    "name": "query_route",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "strategy": {"type": "string", "enum": list(VALID_STRATEGIES)},
            "operation": {"type": ["string", "null"]},
            "filters": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "field": {"type": "string"},
                        "operator": {"type": "string", "enum": ["equals", "contains", "in", "not_equals", "not_in", "gt", "gte", "lt", "lte"]},
                        # `value` is a scalar for equals/contains and an array
                        # for `in`. `anyOf` keeps the Structured Outputs schema
                        # strict while supporting both shapes.
                        "value": {
                            "anyOf": [
                                {"type": "string"},
                                {"type": "array", "items": {"type": "string"}},
                            ]
                        },
                    },
                    "required": ["field", "operator", "value"],
                    "additionalProperties": False,
                },
            },
            "reason": {"type": "string"},
        },
        "required": ["strategy", "operation", "filters", "reason"],
        "additionalProperties": False,
    },
}


def _normalize_strategy(raw):
    """Defense in depth, in case a non-Structured-Outputs fallback path (or a
    future model quirk) still produces something outside the enum — coerce
    it into a sane choice instead of silently defaulting to hybrid_retrieval.
    structured_filter is checked before aggregation so a garbled value like
    "structured_filter | aggregation" resolves to the listing behavior,
    which matches what a "list all ... who ..." question actually wants."""
    if raw in VALID_STRATEGIES:
        return raw
    raw_l = (raw or "").lower()
    for candidate in VALID_STRATEGIES:
        if candidate in raw_l:
            return candidate
    if "aggregat" in raw_l or "count" in raw_l: return "aggregation"
    if "filter" in raw_l: return "structured_filter"
    if "lookup" in raw_l: return "exact_lookup"
    return "hybrid_retrieval"


def _classify_query_prompt(question, fields):
    # NOTE: this is an f-string (interpolates {fields} and {question} below),
    # so every literal JSON brace in the example lines MUST be doubled
    # ({{ / }}) or Python's f-string parser reads it as a real replacement
    # field. A previous version left these examples single-braced — e.g.
    # `-> {"field":"department","operator":"in","value":["Finance","IT"]}` —
    # which crashed EVERY call with "ValueError: Invalid format specifier
    # '"department","operator":"in","value":["Finance","IT"]' for object of
    # type 'str'": Python parsed `"field"` as the expression and everything
    # after the first colon as a (bogus) format spec applied to it.
    return f"""Classify this question.
Available fields: {fields}
Question: {question}
Use structured_filter for lists with field conditions, aggregation for count questions, exact_lookup for direct record lookup, otherwise hybrid_retrieval.
Operators:
- equals: exact single value
- contains: substring matching
- in: one field matches any of multiple values
- not_equals: field does not equal the value
- not_in: field matches none of the listed values
- gt: numeric value greater than the specified value
- gte: numeric value greater than or equal to the specified value
- lt: numeric value less than the specified value
- lte: numeric value less than or equal to the specified value

BOOLEAN RULES:
- "and" means all conditions must match (AND).
- "or" means alternatives must match (OR).
- If multiple values belong to the SAME field and are connected by "or", use one `in` filter with an array of values.
- Example: "Finance or IT" -> {{"field":"department","operator":"in","value":["Finance","IT"]}}
- Do NOT turn "department is Finance or IT" into department=Finance AND department=IT.

NUMERIC COMPARISON RULES:
- "more than", "greater than", "above", "over" -> gt
- "at least", "greater than or equal to", "minimum" -> gte
- "less than", "below", "under" -> lt
- "at most", "less than or equal to", "maximum" -> lte
- "not equal to", "not" -> not_equals when it clearly applies to one value
- Never encode a comparison such as ">50000" as an `in` value.
- Example: "salary more than 50000" -> {{"field":"salary_inr","operator":"gt","value":"50000"}}
- Example: "salary at least 50000" -> {{"field":"salary_inr","operator":"gte","value":"50000"}}
- Example: "salary between 50000 and 75000" -> two filters: salary_inr gte 50000 AND salary_inr lte 75000.

Do not invent fields or values.
IMPORTANT: if the question names a specific value (a status, department, location, salary threshold, or similar), you MUST include a filter for it. Only return an empty filters list when the question truly has no condition at all, such as "how many employees are there in total"."""


def classify_query(question, records):
    fields = sorted({k for r in records for k in r.keys()})
    prompt = _classify_query_prompt(question, fields)
    try:
        response = get_client().chat.completions.create(
            model=CHAT_MODEL, temperature=0,
            response_format={"type": "json_schema", "json_schema": _CLASSIFY_SCHEMA},
            messages=[{"role": "system", "content": "Classify the question into exactly one strategy and, if applicable, structured filters."},
                      {"role": "user", "content": prompt}])
        data = json.loads(response.choices[0].message.content or "{}")
    except Exception:
        # Older models / accounts without Structured Outputs support: fall
        # back to the loose json_object mode.
        try:
            fallback_prompt = prompt + '\nReturn ONLY JSON: {"strategy":"<one of structured_filter, aggregation, exact_lookup, hybrid_retrieval>","operation":null,"filters":[],"reason":"short explanation"}'
            response = get_client().chat.completions.create(
                model=CHAT_MODEL, temperature=0, response_format={"type": "json_object"},
                messages=[{"role": "system", "content": "Return JSON only."}, {"role": "user", "content": fallback_prompt}])
            data = json.loads(response.choices[0].message.content or "{}")
        except Exception:
            return {"strategy": "hybrid_retrieval", "operation": None, "filters": [], "reason": "Router call failed."}
    data["strategy"] = _normalize_strategy(data.get("strategy"))
    data["filters"] = _normalize_filters(question, data.get("filters", []), records)
    return data


def _normalize_filter_value(value):
    """Normalize scalar/list filter values while preserving numeric text."""
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    if value is None:
        return ""
    return str(value).strip()


def _normalize_comparison_filter(operator, value):
    """Recover malformed comparison filters such as in['>50000']."""
    if operator != "in":
        return operator, value
    values = value if isinstance(value, list) else [value]
    if len(values) != 1:
        return operator, value
    raw = str(values[0]).strip()
    for symbol, normalized_operator in ((">=", "gte"), ("<=", "lte"), (">", "gt"), ("<", "lt")):
        if raw.startswith(symbol):
            numeric_value = raw[len(symbol):].strip()
            if numeric_value:
                return normalized_operator, numeric_value
    return operator, value


def _normalize_filters(question, filters, records=None):
    """Normalize router output and recover common malformed boolean/numeric filters."""
    if not isinstance(filters, list):
        return []

    valid_ops = {"equals", "contains", "in", "not_equals", "not_in", "gt", "gte", "lt", "lte"}
    normalized = []

    for raw in filters:
        if not isinstance(raw, dict):
            continue
        field = str(raw.get("field", "")).strip()
        if not field:
            continue
        op = str(raw.get("operator", "equals")).strip().lower()
        if op not in valid_ops:
            op = "equals"
        value = _normalize_filter_value(raw.get("value"))
        op, value = _normalize_comparison_filter(op, value)

        if op in {"in", "not_in"}:
            values = value if isinstance(value, list) else [value]
            values = [v for v in values if v]
            if values:
                normalized.append({"field": field, "operator": op, "value": values})
        else:
            if isinstance(value, list):
                value = value[0] if value else ""
            if value:
                normalized.append({"field": field, "operator": op, "value": value})

    q = (question or "").lower()
    if " or " not in q or len(normalized) < 2:
        return normalized

    # Multiple equals/contains conditions on the SAME field can never all be
    # true for one record at once (a department can't equal both "Finance"
    # and "IT" simultaneously) — whenever the question also contains " or ",
    # that's a reliable signal they were meant as alternatives, not an AND.
    #
    # A previous version tried to *verify* this by locating each value's
    # word position in the question text (via str.find) and checking for a
    # literal " or " between them. That broke on short values like "IT",
    # which are literal substrings of ordinary words ("with", "wait",
    # "unit", ...) — str.find("it") would latch onto the "it" inside "with"
    # long before the real "IT" token, scrambling the position order and
    # silently defeating the merge. This is exactly the bug that turned a
    # Finance/IT "or" question into two separate, mutually-exclusive
    # "contains" filters that could never both match.
    #
    # Grouping by field name alone — no position math at all — fixes this
    # and is simpler: it can't misfire the way substring matching did.
    field_groups: dict[str, list[dict]] = {}
    for condition in normalized:
        if condition["operator"] in {"equals", "contains"}:
            field_groups.setdefault(condition["field"], []).append(condition)

    merged = []
    emitted_fields = set()
    for condition in normalized:
        field = condition["field"]
        if condition["operator"] not in {"equals", "contains"}:
            merged.append(condition)
            continue
        if field in emitted_fields:
            continue
        emitted_fields.add(field)
        group = field_groups[field]
        if len(group) >= 2:
            merged.append({"field": field, "operator": "in", "value": [c["value"] for c in group]})
        else:
            merged.append(group[0])

    return merged


def _field_aliases(field):
    aliases = {field.lower()}
    aliases.update({
        "salary_inr": {"salary", "salary_inr", "pay", "compensation"},
        "department": {"department", "dept"},
        "location": {"location", "city", "office", "place"},
        "status": {"status"},
        "name": {"name", "employee", "employee_name", "full_name"},
    }.get(field, set()))
    return aliases


def _parse_number(value):
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().lower()
    text = text.replace(",", "").replace("₹", "").replace("inr", "").replace("rs.", "").replace("rs", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        raise ValueError(f"Not numeric: {value!r}")
    return float(match.group(0))


def _infer_numeric_filters(question, records):
    q = question.lower()
    numeric_fields = set()
    for r in records:
        for k, v in r.items():
            try:
                _parse_number(v)
                numeric_fields.add(k)
            except (TypeError, ValueError):
                pass

    results = []
    for field in numeric_fields:
        if not any(alias in q for alias in _field_aliases(field)):
            continue

        num = r"(?:₹\s*)?(\d[\d,]*(?:\.\d+)?)\s*(?:inr|rs\.?|rupees?)?"
        between = re.search(rf"\bbetween\s+{num}\s+and\s+{num}", q)
        if between:
            results.extend([
                {"field": field, "operator": "gte", "value": between.group(1).replace(",", "")},
                {"field": field, "operator": "lte", "value": between.group(2).replace(",", "")},
            ])
            continue

        patterns = [
            (r"(?:more than|greater than|above|over)\s+" + num, "gt"),
            (r"(?:at least|greater than or equal to|minimum(?: of)?)\s+" + num, "gte"),
            (r"(?:less than|below|under)\s+" + num, "lt"),
            (r"(?:at most|less than or equal to|maximum(?: of)?)\s+" + num, "lte"),
        ]
        for pattern, op in patterns:
            match = re.search(pattern, q)
            if match:
                results.append({"field": field, "operator": op, "value": match.group(1).replace(",", "")})
                break
    return results


def infer_filters_from_question(question, records):
    """Safety net for dropped router filters, including numeric comparisons.

    Field values in structured data are a closed vocabulary, so known values
    can be detected directly in the question text, independent of whatever
    the LLM router did or didn't return. Matching uses word boundaries
    rather than a plain substring search: a short value like "IT" or "US" is
    a literal substring of ordinary English words ("with", "wait", "unit",
    ...), so a naive `in`/`str.find` check can silently match the wrong part
    of the question — this is exactly what broke the previous version's
    position-based "or" detection for a "Finance or IT" style question.
    """
    if not records:
        return []

    numeric = _infer_numeric_filters(question, records)
    numeric_fields = {f["field"] for f in numeric}

    field_values = {}
    for r in records:
        for k, v in r.items():
            if isinstance(v, str) and v.strip():
                field_values.setdefault(k, set()).add(v.strip())

    q_lower = question.lower()
    has_or = " or " in q_lower

    candidates = []  # (value_len, field, value) — sorted so longer matches win ties
    for field, values in field_values.items():
        if field in numeric_fields:
            continue
        for value in values:
            if len(value) < 2:
                continue
            pattern = r"(?<!\w)" + re.escape(value.lower()) + r"(?!\w)"
            if re.search(pattern, q_lower):
                candidates.append((len(value), field, value))

    candidates.sort(reverse=True)

    order = []
    grouped: dict[str, list[str]] = {}
    for _, field, value in candidates:
        if field not in grouped:
            grouped[field] = []
            order.append(field)
        grouped[field].append(value)

    inferred = list(numeric)
    for field in order:
        values = grouped[field]
        if len(values) >= 2 and has_or:
            inferred.append({"field": field, "operator": "in", "value": list(dict.fromkeys(values))})
        else:
            inferred.append({"field": field, "operator": "equals", "value": values[0]})

    return inferred


def apply_filters(records, filters):
    """Apply structured filters. Separate filter objects are AND-ed together."""
    normalized_filters = _normalize_filters("", filters)
    result = []

    for record in records:
        ok = True
        for condition in normalized_filters:
            actual = record.get(condition.get("field"))
            expected = condition.get("value")
            op = condition.get("operator", "equals")

            if actual is None:
                ok = False
                break

            if op in {"gt", "gte", "lt", "lte"}:
                try:
                    a_num = _parse_number(actual)
                    e_num = _parse_number(expected)
                except (TypeError, ValueError):
                    ok = False
                    break
                if op == "gt" and not a_num > e_num:
                    ok = False
                elif op == "gte" and not a_num >= e_num:
                    ok = False
                elif op == "lt" and not a_num < e_num:
                    ok = False
                elif op == "lte" and not a_num <= e_num:
                    ok = False
                if not ok:
                    break
                continue

            a = str(actual).strip().lower()
            if op == "equals":
                if a != str(expected).strip().lower():
                    ok = False
                    break
            elif op == "contains":
                if str(expected).strip().lower() not in a:
                    ok = False
                    break
            elif op == "in":
                values = expected if isinstance(expected, list) else [expected]
                if a not in {str(v).strip().lower() for v in values}:
                    ok = False
                    break
            elif op == "not_equals":
                if a == str(expected).strip().lower():
                    ok = False
                    break
            elif op == "not_in":
                values = expected if isinstance(expected, list) else [expected]
                if a in {str(v).strip().lower() for v in values}:
                    ok = False
                    break

        if ok:
            result.append(record)
    return result


def exact_lookup(question, records, filters=None):
    """Look up a specific record. The router's filters (e.g.
    {"field":"name","operator":"equals","value":"Vikram Malhotra"}) are the
    reliable path — they target one field directly, the same way
    apply_filters() works for structured_filter. The old behavior checked
    whether the ENTIRE question string was a substring of a single field's
    value (`q in str(v).lower()`), which only ever matches if you type the
    exact field value as your whole question (e.g. just "EMP001") — any
    natural-language question like "who is the manager of Vikram Malhotra?"
    can never be a substring of a "Vikram Malhotra" field value, so it
    always returned zero records despite the router correctly identifying
    who to look up. Filters are tried first now; the substring heuristic is
    kept only as a fallback for when the router doesn't supply filters."""
    if filters:
        matched = apply_filters(records, filters)
        if matched:
            return matched
    q = question.lower()
    return [r for r in records if any(q in str(v).lower() for v in r.values() if v is not None)]


def hybrid_search(question, document_ids, top_k):
    """document_ids: None to search across every indexed document, or a list
    of document_id strings to scope the search to those documents only."""
    if document_ids is None:
        data = _collection.get(include=["documents", "metadatas", "embeddings"])
    else:
        data = _collection.get(where={"document_id": {"$in": document_ids}}, include=["documents", "metadatas", "embeddings"])
    chunks, metadatas = data["documents"], data["metadatas"]
    if not chunks: return []
    embeddings = np.array(data["embeddings"], dtype=np.float32)
    q = embed([question])[0]; q = q / (np.linalg.norm(q) + 1e-8)
    d = embeddings / (np.linalg.norm(embeddings, axis=1, keepdims=True) + 1e-8)
    vector_rank = np.argsort(d @ q)[::-1].tolist()
    bm25 = BM25Okapi([tokenize(x) for x in chunks]); bm25_rank = np.argsort(bm25.get_scores(tokenize(question)))[::-1].tolist()
    fused = {}
    for ranking in (vector_rank, bm25_rank):
        for rank, idx in enumerate(ranking): fused[idx] = fused.get(idx, 0) + 1 / (RRF_K + rank + 1)
    final = sorted(fused, key=fused.get, reverse=True)[:top_k]
    return [{"rank": rank, "chunk_index": metadatas[idx]["chunk_index"], "filename": metadatas[idx]["filename"], "text": chunks[idx],
             "vector_rank": vector_rank.index(idx) + 1, "bm25_rank": bm25_rank.index(idx) + 1}
            for rank, idx in enumerate(final, 1)]


def generate_answer(question, evidence, strategy):
    prompt = f"""Answer ONLY from the verified evidence.
Question: {question}
Strategy: {strategy}
Evidence:
{json.dumps(evidence, ensure_ascii=False, indent=2)}
Rules: do not invent facts; use supplied structured records for lists; if the evidence is empty or insufficient, say so plainly instead of guessing; keep concise."""
    response = get_client().chat.completions.create(model=CHAT_MODEL, temperature=0, messages=[
        {"role": "system", "content": "You are a grounded RAG assistant."}, {"role": "user", "content": prompt}])
    return response.choices[0].message.content or ""


def _describe_filters(filters):
    if not filters:
        return None
    parts = []
    for f in filters:
        field = f.get("field", "?")
        op = f.get("operator", "equals")
        value = f.get("value", "?")
        if op == "contains":
            parts.append(f'{field} contains "{value}"')
        elif op in {"in", "not_in"}:
            values = value if isinstance(value, list) else [value]
            rendered = ", ".join(f'"{v}"' for v in values)
            parts.append(f"{field} {op.replace('_', ' ')} [{rendered}]")
        elif op == "not_equals":
            parts.append(f'{field} != "{value}"')
        elif op in {"gt", "gte", "lt", "lte"}:
            symbol = {"gt": ">", "gte": ">=", "lt": "<", "lte": "<="}[op]
            parts.append(f'{field} {symbol} {value}')
        else:
            parts.append(f'{field} = "{value}"')
    return " and ".join(parts)


def generate_aggregation_answer(question, count, filters, matched_records):
    """Aggregation answers must never let the LLM recount a list of records.
    LLMs are unreliable at counting items enumerated in their own context —
    handing over a full list of e.g. 23 matched records and asking "how many
    are there" invites the model to recount from the list itself and get it
    wrong, even when the correct count is already given to it. len(matched)
    is computed here in plain Python and is authoritative; the LLM is only
    asked to phrase that number into a sentence, never to derive it. If the
    model's phrasing ever fails to include the authoritative number anyway,
    we discard its prose and fall back to a deterministic sentence."""
    filter_desc = _describe_filters(filters)
    prompt = f"""Question: {question}
The exact, verified count is: {count}. This number is already correct and final — it was computed directly from the data, not by you.
Filters applied: {filter_desc or "none"}
Sample of matching records (for context only, NOT for recounting): {json.dumps(matched_records[:3], ensure_ascii=False)}
Write one concise sentence answering the question. You MUST state the number {count} exactly as given. Do not recount records, do not adjust the number, do not estimate."""
    response = get_client().chat.completions.create(model=CHAT_MODEL, temperature=0, messages=[
        {"role": "system", "content": "You are a grounded RAG assistant. A count has already been computed for you and is never to be second-guessed or recalculated."},
        {"role": "user", "content": prompt}])
    text = response.choices[0].message.content or ""
    if re.search(rf"(?<!\d){count}(?!\d)", text) is None:
        suffix = f" matching {filter_desc}" if filter_desc else ""
        text = f"There are {count} matching record(s){suffix}."
    return text


def query_document(question, document_id=None, top_k=TOP_K):
    """document_id=None queries across every indexed document; pass a
    specific document_id to scope the question to just that one file."""
    all_documents = list_documents()
    if not all_documents: raise RuntimeError("No documents are indexed.")
    if document_id:
        scoped_docs = [d for d in all_documents if d["document_id"] == document_id]
        if not scoped_docs: raise RuntimeError("Selected document was not found (it may have been deleted).")
        doc_ids = [document_id]
    else:
        scoped_docs = all_documents
        doc_ids = None

    records = []
    for d in scoped_docs: records.extend(get_records(d["document_id"]))
    document_label = scoped_docs[0]["filename"] if len(scoped_docs) == 1 else f"All documents ({len(scoped_docs)})"

    route = classify_query(question, records); strategy = route.get("strategy", "hybrid_retrieval"); filters = _normalize_filters(question, route.get("filters", []), records)
    route = {**route, "filters": filters}
    if strategy == "hybrid_retrieval" and filters:
        # Extra insurance: the prompt only ever populates filters for
        # structured strategies, so a non-empty filter list alongside a
        # hybrid_retrieval label means the strategy label itself is the
        # unreliable part, not the parsed conditions. Trust the filters —
        # this is exactly the failure mode observed where the model
        # returned "structured_filter | aggregation" as the strategy string,
        # got normalized away, but the well-formed filters were still there
        # and would otherwise have been discarded.
        strategy = "aggregation" if route.get("operation") == "count" else "structured_filter"
    if strategy in ("structured_filter", "aggregation", "exact_lookup") and not filters:
        # The router said this needed filtering but returned none — an empty
        # filter list matches every record, so trust it only after checking
        # whether the question actually names a known value we can detect
        # ourselves. See infer_filters_from_question() for why this matters.
        inferred = infer_filters_from_question(question, records)
        if inferred:
            filters = _normalize_filters(question, inferred, records)
            route = {**route, "filters": filters, "reason": route.get("reason", "") + " (filters auto-detected from question text; router returned none)"}
    if strategy == "structured_filter":
        matched = apply_filters(records, filters)
        return {"answer": generate_answer(question, matched, "Structured filtering"), "strategy": "Structured filtering", "route": route, "filters": filters, "matched_records": matched, "record_count": len(matched), "sources": [], "document_name": document_label}
    if strategy == "aggregation":
        matched = apply_filters(records, filters)
        answer = generate_aggregation_answer(question, len(matched), filters, matched)
        return {"answer": answer, "strategy": "Structured aggregation", "route": route, "filters": filters, "matched_records": matched, "record_count": len(matched), "sources": [], "document_name": document_label}
    if strategy == "exact_lookup":
        matched = exact_lookup(question, records, filters)
        return {"answer": generate_answer(question, matched, "Exact structured lookup"), "strategy": "Exact structured lookup", "route": route, "filters": filters, "matched_records": matched, "record_count": len(matched), "sources": [], "document_name": document_label}
    hits = hybrid_search(question, doc_ids, top_k)
    return {"answer": generate_answer(question, [{"rank": h["rank"], "filename": h["filename"], "text": h["text"]} for h in hits], "Hybrid retrieval"),
            "strategy": "Hybrid retrieval", "route": route, "filters": [], "matched_records": [], "record_count": 0, "sources": hits, "document_name": document_label}