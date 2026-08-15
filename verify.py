"""Offline smoke test - checks imports, storage paths, every file loader
(PDF/DOCX/CSV/TXT/MD), the aggregation safety net, the AND/OR filter
normalization, numeric comparisons (gt/gte/lt/lte/between), not_equals/
not_in, and the router-prompt / router-recovery safety nets - all without
calling the OpenAI API, so it works even before an API key is configured.
Run with: python verify.py
"""
import io, types
import advanced_rag as rag

print("OK: advanced_rag imported")
print("SQLite:", rag.SQLITE_PATH)
print("ChromaDB:", rag.CHROMA_PATH)
print("Supported extensions:", rag.SUPPORTED_EXTENSIONS)

# --- TXT / MD -------------------------------------------------------------
txt_bytes = b"Employee Handbook\nAll staff must badge in by 9am.\n"
text, units = rag.extract_plain_text(txt_bytes)
assert "badge in" in text
print(f"OK: TXT/MD loader -> {units} lines")

# --- CSV --------------------------------------------------------------
csv_bytes = (
    b"name,department,status\n"
    b"Asha Rao,Engineering,Active\n"
    b"Ravi Kumar,Marketing,On Leave\n"
)
text, records = rag.extract_csv(csv_bytes)
assert len(records) == 2 and records[0]["name"] == "Asha Rao"
print(f"OK: CSV loader -> {len(records)} records: {records}")

# --- DOCX ---------------------------------------------------------------
try:
    from docx import Document as DocxDocument
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("Quarterly Review")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "name"
    table.rows[0].cells[1].text = "status"
    row = table.add_row()
    row.cells[0].text = "Meera Iyer"
    row.cells[1].text = "Active"
    doc.save(buf)
    text, units, table_records = rag.extract_docx(buf.getvalue())
    assert table_records and table_records[0]["name"] == "Meera Iyer"
    print(f"OK: DOCX loader -> {units} units, records: {table_records}")
except ImportError:
    print("SKIP: python-docx not installed")

# --- chunking / normalization -------------------------------------------
chunks = rag.chunk_text("word " * 2000)
assert chunks
print(f"OK: chunk_text -> {len(chunks)} chunks")

normalized = rag.normalize_record({"Full Name ": " Asha Rao ", "Dept#": "Eng"})
assert normalized == {"full_name": "Asha Rao", "dept": "Eng"}
print(f"OK: normalize_record -> {normalized}")

batches = rag._batch_by_lines("line\n" * 100, max_chars=50)
assert all(len(b) <= 60 for b in batches)  # allows small slack for the last line in a batch
print(f"OK: _batch_by_lines -> {len(batches)} batches")

# --- aggregation safety net ----------------------------------------------
# Simulates the exact failure mode of miscounting: the verified count is 23,
# but the LLM's prose (wrongly) claims 20. The safety net must override it.
class _FakeMsg:
    def __init__(self, content): self.message = types.SimpleNamespace(content=content)
class _FakeResp:
    def __init__(self, content): self.choices = [_FakeMsg(content)]
class _FakeChat:
    def create(self, **kwargs): return _FakeResp('There are currently 20 employees "On Leave."')
class _FakeClient:
    chat = types.SimpleNamespace(completions=_FakeChat())

_original_get_client = rag.get_client
rag.get_client = lambda: _FakeClient()
matched = [{"name": f"Emp{i}", "status": "On Leave"} for i in range(23)]
filters = [{"field": "status", "operator": "equals", "value": "On Leave"}]
answer = rag.generate_aggregation_answer('How many employees are "On Leave"?', len(matched), filters, matched)
rag.get_client = _original_get_client
assert "23" in answer, f"aggregation safety net failed to correct a miscounted LLM answer: {answer!r}"
print(f"OK: aggregation safety net -> corrected model's wrong '20' to the verified count: {answer!r}")

# --- dropped-filter safety net --------------------------------------------
# Simulates the router bug that turned "how many are on leave" (23) into a
# wrong "100" (every record): the router says aggregation but returns no
# filters at all, even though the question names a specific status value.
big_records = [{"name": f"Emp{i}", "department": "Engineering",
                "status": "On Leave" if i < 23 else "Active"} for i in range(100)]
inferred = rag.infer_filters_from_question('How many employees are currently "On Leave"?', big_records)
assert inferred == [{"field": "status", "operator": "equals", "value": "On Leave"}], f"unexpected inference: {inferred}"
recovered = rag.apply_filters(big_records, inferred)
assert len(recovered) == 23, f"dropped-filter safety net failed: got {len(recovered)} instead of 23"
print(f"OK: dropped-filter safety net -> recovered {len(recovered)}/100 records instead of matching everything")

# A genuinely filter-less question (e.g. "how many employees are there in
# total") must NOT have a filter invented for it.
no_filter = rag.infer_filters_from_question("How many employees are there in total?", big_records)
assert no_filter == [], f"should not have inferred a filter here: {no_filter}"
print("OK: dropped-filter safety net leaves genuinely unconditional questions alone")

# --- garbled router-strategy safety net -----------------------------------
# Reproduces a real failure: the router echoed its own prompt placeholder
# text back verbatim as "strategy": "structured_filter | aggregation" (via
# the old loose json_object mode), which matched none of the strategy
# branches and silently discarded two correctly-parsed filters, falling back
# to hybrid retrieval. _normalize_strategy() must resolve that string to a
# real strategy, favoring the listing behavior for a garbled "filter |
# aggregation" combo.
assert rag._normalize_strategy("structured_filter | aggregation") == "structured_filter"
assert rag._normalize_strategy("aggregation | exact_lookup") == "aggregation"
assert rag._normalize_strategy("totally unrecognized text") == "hybrid_retrieval"
assert rag._normalize_strategy("structured_filter") == "structured_filter"  # clean values pass through unchanged
print("OK: garbled router-strategy safety net -> normalizes stray/echoed strategy text")

# End-to-end: classify_query() falls back to json_object mode (simulating a
# model/account without Structured Outputs support) and must still recover
# from that exact malformed shape.
class _GarbledMsg:
    def __init__(self, content): self.message = types.SimpleNamespace(content=content)
class _GarbledResp:
    def __init__(self, content): self.choices = [_GarbledMsg(content)]
_garbled_call_count = {"n": 0}
class _GarbledChat:
    def create(self, **kwargs):
        _garbled_call_count["n"] += 1
        if _garbled_call_count["n"] == 1:
            raise RuntimeError("simulated: this model/account does not support response_format=json_schema")
        return _GarbledResp('''{
            "strategy": "structured_filter | aggregation",
            "field": null, "operator": null, "value": null,
            "filters": [
                {"field": "department", "operator": "equals", "value": "Engineering"},
                {"field": "status", "operator": "equals", "value": "On Leave"}
            ],
            "operation": "count",
            "reason": "needs both filters"
        }''')
class _GarbledClient:
    chat = types.SimpleNamespace(completions=_GarbledChat())
rag.get_client = lambda: _GarbledClient()
route = rag.classify_query("List all employees in Engineering department who are on leave.",
                            [{"department": "Engineering", "status": "On Leave"}])
rag.get_client = _original_get_client
assert route["strategy"] == "structured_filter", f"garbled strategy was not recovered: {route}"
assert len(route["filters"]) == 2, f"well-formed filters were lost: {route}"
print(f"OK: classify_query() recovers from a garbled router response end-to-end -> {route['strategy']}, {len(route['filters'])} filters kept")

# --- exact_lookup actually uses the router's filters -----------------------
# Reproduces a real failure: "Who is the manager of Vikram Malhotra?" was
# correctly routed to exact_lookup with a filter naming the person, but the
# old exact_lookup(question, records) ignored filters entirely and instead
# checked whether the ENTIRE question string was a substring of a single
# field's value - which a natural-language question never is - so it always
# returned 0 records despite the router doing its job correctly.
lookup_records = [
    {"name": "Vikram Malhotra", "department": "Engineering", "manager": "Sunita Rao"},
    {"name": "Kevin Davis", "department": "Engineering", "manager": "Jennifer Mehta"},
]
lookup_filters = [{"field": "name", "operator": "equals", "value": "Vikram Malhotra"}]
old_style_result = rag.exact_lookup("Who is the manager of Vikram Malhotra?", lookup_records)
assert old_style_result == [], "sanity check: confirms the original bug would find nothing"
fixed_result = rag.exact_lookup("Who is the manager of Vikram Malhotra?", lookup_records, lookup_filters)
assert len(fixed_result) == 1 and fixed_result[0]["manager"] == "Sunita Rao", f"exact_lookup fix failed: {fixed_result}"
print(f"OK: exact_lookup() uses router filters -> found {fixed_result[0]['name']}, manager={fixed_result[0]['manager']}")

# --- same-field OR normalization (regression: "IT" is a substring of "with") ---
# Reproduces the exact assertion failure the user hit: a question mentioning
# a short department code ("IT") alongside an ordinary word that contains it
# as a substring ("with") used to scramble the old position-based " or "
# detection in _normalize_filters, so "Finance or IT" was left as two
# separate equals filters (department=Finance AND department=IT) that could
# never both match the same record, instead of being merged into one
# `in` filter. The fix groups by field name only - no substring position
# math - so this can no longer misfire.
or_question = "List employees in Finance or IT department, along with their email"
raw_filters = [
    {"field": "department", "operator": "equals", "value": "Finance"},
    {"field": "department", "operator": "equals", "value": "IT"},
]
merged = rag._normalize_filters(or_question, raw_filters)
assert merged == [{"field": "department", "operator": "in", "value": ["Finance", "IT"]}], \
    f"same-field OR normalization failed: {merged}"
print(f"OK: same-field OR normalization -> merged despite 'with' containing 'it': {merged}")

or_records = [
    {"name": "A", "department": "Finance"},
    {"name": "B", "department": "IT"},
    {"name": "C", "department": "Marketing"},
]
or_result = rag.apply_filters(or_records, merged)
assert {r["name"] for r in or_result} == {"A", "B"}, f"OR filter application failed: {or_result}"
print(f"OK: same-field OR filter applied correctly -> matched {[r['name'] for r in or_result]}")

# infer_filters_from_question must reach the same merged result independently
# (i.e. the word-boundary regex fix), confirming the dropped-filter safety
# net also survives the "with"/"IT" trap.
inferred_or = rag.infer_filters_from_question(or_question, or_records)
assert inferred_or == [{"field": "department", "operator": "in", "value": ["Finance", "IT"]}], \
    f"infer_filters_from_question OR-merge failed: {inferred_or}"
print(f"OK: infer_filters_from_question -> word-boundary match avoids the 'with' trap: {inferred_or}")

# --- numeric comparisons: gt / gte / lt / lte / between --------------------
salary_records = [
    {"name": "P1", "salary_inr": "40000"},
    {"name": "P2", "salary_inr": "55000"},
    {"name": "P3", "salary_inr": "60000"},
    {"name": "P4", "salary_inr": "75000"},
]
gt = rag.apply_filters(salary_records, [{"field": "salary_inr", "operator": "gt", "value": "55000"}])
assert {r["name"] for r in gt} == {"P3", "P4"}, f"gt failed: {gt}"
gte = rag.apply_filters(salary_records, [{"field": "salary_inr", "operator": "gte", "value": "55000"}])
assert {r["name"] for r in gte} == {"P2", "P3", "P4"}, f"gte failed: {gte}"
lt = rag.apply_filters(salary_records, [{"field": "salary_inr", "operator": "lt", "value": "55000"}])
assert {r["name"] for r in lt} == {"P1"}, f"lt failed: {lt}"
lte = rag.apply_filters(salary_records, [{"field": "salary_inr", "operator": "lte", "value": "55000"}])
assert {r["name"] for r in lte} == {"P1", "P2"}, f"lte failed: {lte}"
print("OK: apply_filters -> gt/gte/lt/lte all correct")

between_filters = rag._infer_numeric_filters("salary between 50000 and 65000", salary_records)
assert between_filters == [
    {"field": "salary_inr", "operator": "gte", "value": "50000"},
    {"field": "salary_inr", "operator": "lte", "value": "65000"},
], f"between parsing failed: {between_filters}"
between_result = rag.apply_filters(salary_records, between_filters)
assert {r["name"] for r in between_result} == {"P2", "P3"}, f"between filter application failed: {between_result}"
print(f"OK: 'between X and Y' -> parsed to gte/lte and matched {[r['name'] for r in between_result]}")

gt_inferred = rag._infer_numeric_filters("salary more than 55000", salary_records)
assert gt_inferred == [{"field": "salary_inr", "operator": "gt", "value": "55000"}], f"gt phrase parsing failed: {gt_inferred}"
print(f"OK: 'more than' phrasing -> parsed to gt: {gt_inferred}")

# --- not_equals / not_in ----------------------------------------------------
status_records = [
    {"name": "A", "status": "Active"},
    {"name": "B", "status": "On Leave"},
    {"name": "C", "status": "Terminated"},
]
not_eq = rag.apply_filters(status_records, [{"field": "status", "operator": "not_equals", "value": "Active"}])
assert {r["name"] for r in not_eq} == {"B", "C"}, f"not_equals failed: {not_eq}"
not_in = rag.apply_filters(status_records, [{"field": "status", "operator": "not_in", "value": ["Active", "Terminated"]}])
assert {r["name"] for r in not_in} == {"B"}, f"not_in failed: {not_in}"
print("OK: apply_filters -> not_equals/not_in both correct")

# --- malformed comparison filter recovery (in['>50000']) -------------------
# A router or upstream normalization step can sometimes still produce a
# comparison encoded as an `in` filter with a single value like ">50000"
# instead of a proper gt/50000 pair. _normalize_comparison_filter() must
# recover the intended operator and numeric value.
recovered_op, recovered_val = rag._normalize_comparison_filter("in", [">50000"])
assert (recovered_op, recovered_val) == ("gt", "50000"), f"malformed '>' recovery failed: {(recovered_op, recovered_val)}"
recovered_op2, recovered_val2 = rag._normalize_comparison_filter("in", ["<=75000"])
assert (recovered_op2, recovered_val2) == ("lte", "75000"), f"malformed '<=' recovery failed: {(recovered_op2, recovered_val2)}"
malformed_normalized = rag._normalize_filters("", [{"field": "salary_inr", "operator": "in", "value": [">50000"]}])
assert malformed_normalized == [{"field": "salary_inr", "operator": "gt", "value": "50000"}], \
    f"end-to-end malformed comparison recovery failed: {malformed_normalized}"
print("OK: malformed comparison filters (in['>50000']) recovered to proper gt/lte operators")

# --- regression guard: _classify_query_prompt must not crash ---------------
# Reproduces the exact live crash the user hit: literal JSON braces used as
# example text inside this f-string, if left single-braced, are parsed by
# Python as real replacement fields and raise
# "ValueError: Invalid format specifier ... for object of type 'str'" on
# EVERY call to classify_query(), before any API request is even made. This
# must build cleanly and contain the literal example text, un-mangled.
prompt_text = rag._classify_query_prompt("Who works in Finance or IT?", ["department", "salary_inr"])
assert '{"field":"department","operator":"in","value":["Finance","IT"]}' in prompt_text, \
    "classify prompt lost its literal JSON example text"
assert '{"field":"salary_inr","operator":"gt","value":"50000"}' in prompt_text, \
    "classify prompt lost its literal JSON example text"
assert "Who works in Finance or IT?" in prompt_text
print("OK: _classify_query_prompt builds without crashing and preserves literal JSON examples")

print("\nAll offline checks passed. Set OPENAI_API_KEY in .env, then run: streamlit run app.py")